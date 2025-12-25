from docx import Document
from io import BytesIO

def extract_from_doc(file_content: bytes) -> str:
    """
    Extract text from DOCX file including paragraphs, tables, and lists.
    
    Args:
        file_content: DOCX file as bytes
        
    Returns:
        Extracted text as string
        
    Raises:
        ValueError: If DOCX parsing fails
    """
    try:
        doc = Document(BytesIO(file_content))
        
        all_text = []
        
        # Extract paragraphs (including bullet points)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(text)
        
        # IMPORTANT: Extract tables (skills often in tables!)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    # Join cells with space or comma
                    all_text.append(" | ".join(row_text))
        
        # Join all extracted text with double line breaks
        full_text = "\n\n".join(all_text)
        return full_text.strip()
        
    except Exception as e:
        raise ValueError(f"Error parsing DOCX: {str(e)}")



